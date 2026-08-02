"""Publishable manuscript exports from a Project."""

from __future__ import annotations

import html
import io
import json
import re
import zipfile
from xml.sax.saxutils import escape as xml_escape

from app.models.schemas import Project


def _slug(title: str) -> str:
    s = re.sub(r"[^\w\s-]", "", title or "manuscript", flags=re.U)
    s = re.sub(r"[-\s]+", "-", s.strip()).strip("-").lower()
    return s or "manuscript"


def _sorted_chapters(project: Project):
    return sorted(project.chapters or [], key=lambda c: c.order)


_SCENE_BREAK_RE = re.compile(
    r"^(?:\*\s*\*\s*\*|\*\*\*|---|—{2,}|#{3,}|·\s*·\s*·)$"
)


def _paragraphs(text: str) -> list[str]:
    """Plain paragraph strings (no scene-break detection)."""
    return [body for kind, body in _novel_blocks(text) if kind == "p"]


def _novel_blocks(text: str) -> list[tuple[str, str]]:
    """
    Split manuscript text into novel blocks.

    Returns list of (kind, text) where kind is 'p' | 'break'.
    Single newlines inside a paragraph become spaces (ebook-friendly).
    """
    if not (text or "").strip():
        return []
    chunks = re.split(r"\n\s*\n", text.strip())
    out: list[tuple[str, str]] = []
    for raw in chunks:
        block = raw.strip()
        if not block:
            continue
        # Normalize internal newlines → spaces (avoids jagged ebook lines)
        flat = re.sub(r"[ \t]*\n[ \t]*", " ", block)
        flat = re.sub(r" {2,}", " ", flat).strip()
        if _SCENE_BREAK_RE.match(flat):
            out.append(("break", ""))
        else:
            out.append(("p", flat))
    return out


def _chapter_heading_parts(title: str, index: int) -> tuple[str, str | None]:
    """
    Split 'Chapter 3: The Bus' into (kicker, subtitle).
    """
    t = (title or f"Chapter {index}").strip()
    m = re.match(
        r"^(chapter\s+(\d+|[ivxlcdm]+))\s*[:.\-—–]\s*(.+)$",
        t,
        flags=re.I,
    )
    if m:
        return m.group(1).strip(), m.group(3).strip()
    m2 = re.match(r"^(chapter\s+(\d+|[ivxlcdm]+))$", t, flags=re.I)
    if m2:
        return m2.group(1).strip(), None
    # Untitled-style: use Chapter N as kicker, title as subtitle
    if not re.match(r"^chapter\b", t, flags=re.I):
        return f"Chapter {index}", t
    return t, None


def _xhtml_escape_text(s: str) -> str:
    return xml_escape(s, {"\"": "&quot;", "'": "&apos;"})


def filename_for(project: Project, fmt: str) -> str:
    base = _slug(project.title)
    ext = {
        "markdown": "md",
        "md": "md",
        "txt": "txt",
        "text": "txt",
        "html": "html",
        "docx": "docx",
        "epub": "epub",
        "json": "json",
        "manuscript-docx": "docx",
        "manuscript-epub": "epub",
        "cover-jpg": "jpg",
        "cover-tiff": "tiff",
    }.get(fmt, fmt)
    if fmt in ("manuscript-docx", "manuscript-epub"):
        return f"{base}-manuscript.{ext}"
    if fmt in ("cover-jpg", "cover-tiff"):
        return f"{base}-cover.{ext}"
    return f"{base}.{ext}"


def media_type_for(fmt: str) -> str:
    return {
        "markdown": "text/markdown; charset=utf-8",
        "md": "text/markdown; charset=utf-8",
        "txt": "text/plain; charset=utf-8",
        "text": "text/plain; charset=utf-8",
        "html": "text/html; charset=utf-8",
        "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "manuscript-docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "epub": "application/epub+zip",
        "manuscript-epub": "application/epub+zip",
        "json": "application/json; charset=utf-8",
        "cover-jpg": "image/jpeg",
        "cover-tiff": "image/tiff",
    }.get(fmt, "application/octet-stream")


def export_project(project: Project, fmt: str) -> bytes:
    fmt = (fmt or "markdown").lower().strip()
    if fmt in ("markdown", "md"):
        return _to_markdown(project).encode("utf-8")
    if fmt in ("txt", "text"):
        return _to_text(project).encode("utf-8")
    if fmt == "html":
        return _to_html(project).encode("utf-8")
    if fmt == "docx":
        return _to_docx(project)
    if fmt == "manuscript-docx":
        return _to_docx_manuscript(project)
    if fmt == "epub":
        return _to_epub(project)
    if fmt == "manuscript-epub":
        return _to_epub_manuscript(project)
    if fmt in ("cover-jpg", "cover-tiff"):
        return _to_cover_image(project, fmt)
    if fmt == "json":
        return (
            json.dumps(project.model_dump(), indent=2, ensure_ascii=False) + "\n"
        ).encode("utf-8")
    raise ValueError(f"Unsupported export format: {fmt}")


def _front_matter_lines(project: Project) -> list[str]:
    lines = [project.title or "Untitled"]
    if project.premise:
        lines.append(project.premise)
    elif project.description:
        lines.append(project.description)
    if project.genre:
        lines.append(project.genre)
    return lines


def _to_markdown(project: Project) -> str:
    parts: list[str] = [f"# {project.title or 'Untitled'}", ""]
    if project.premise:
        parts += [f"*{project.premise}*", ""]
    elif project.description:
        parts += [f"*{project.description}*", ""]
    if project.genre:
        parts += [f"**Genre:** {project.genre}", ""]

    if project.characters:
        parts += ["## Dramatis Personae", ""]
        for c in project.characters:
            role = f" — {c.role}" if c.role else ""
            parts.append(f"- **{c.name}**{role}")
        parts.append("")

    for i, ch in enumerate(_sorted_chapters(project), start=1):
        title = ch.title or f"Chapter {i}"
        parts += [f"## {title}", ""]
        body = (ch.content or "").strip()
        if body:
            parts.append(body)
        else:
            parts.append("*(empty)*")
        parts += ["", ""]

    if (project.world_notes or "").strip():
        parts += ["## World Notes", "", project.world_notes.strip(), ""]

    return "\n".join(parts).rstrip() + "\n"


def _to_text(project: Project) -> str:
    parts: list[str] = []
    for line in _front_matter_lines(project):
        parts.append(line)
    parts += ["", "=" * 40, ""]

    for i, ch in enumerate(_sorted_chapters(project), start=1):
        title = ch.title or f"Chapter {i}"
        parts += [title, "-" * len(title), "", (ch.content or "").strip(), "", ""]

    return "\n".join(parts).rstrip() + "\n"


def _html_body_from_blocks(blocks: list[tuple[str, str]], *, first_nofirst: bool = True) -> str:
    parts: list[str] = []
    first_p = True
    for kind, text in blocks:
        if kind == "break":
            parts.append('<p class="scenebreak">* * *</p>')
            first_p = True
            continue
        cls = ' class="first"' if first_p and first_nofirst else ""
        parts.append(f"<p{cls}>{html.escape(text)}</p>")
        first_p = False
    return "".join(parts) if parts else "<p class=\"first\"><em>(empty)</em></p>"


def _to_html(project: Project) -> str:
    title = html.escape(project.title or "Untitled")
    meta_bits = []
    if project.premise:
        meta_bits.append(
            f"<p class='premise'><em>{html.escape(project.premise)}</em></p>"
        )
    if project.genre:
        meta_bits.append(f"<p class='genre'>{html.escape(project.genre)}</p>")

    chapters_html = []
    for i, ch in enumerate(_sorted_chapters(project), start=1):
        kicker, sub = _chapter_heading_parts(ch.title or f"Chapter {i}", i)
        if sub:
            head = (
                f"<p class='ch-kicker'>{html.escape(kicker)}</p>"
                f"<h2 class='ch-title'>{html.escape(sub)}</h2>"
            )
        else:
            head = f"<h2 class='ch-title'>{html.escape(kicker)}</h2>"
        body = _html_body_from_blocks(_novel_blocks(ch.content or ""))
        chapters_html.append(
            f"<section class='chapter' id='ch-{i}'>\n{head}\n{body}\n</section>"
        )

    return f"""<!DOCTYPE html>
<html lang="en">
<head>
  <meta charset="utf-8"/>
  <meta name="viewport" content="width=device-width, initial-scale=1"/>
  <title>{title}</title>
  <style>
    @page {{ size: 5.5in 8.5in; margin: 0.75in 0.7in; }}
    body {{
      max-width: 32rem; margin: 2rem auto; padding: 0 1.25rem 4rem;
      font-family: "Palatino Linotype", Palatino, "Book Antiqua", Georgia, serif;
      font-size: 12pt; line-height: 1.55; color: #1a1814;
      background: #faf8f2;
      hyphens: auto; -webkit-hyphens: auto;
    }}
    header.titlepage {{
      min-height: 70vh; display: flex; flex-direction: column;
      justify-content: center; align-items: center; text-align: center;
      page-break-after: always;
    }}
    h1.title {{
      font-weight: 600; font-size: 2rem; letter-spacing: 0.04em;
      margin: 0 0 1.5rem; text-wrap: balance;
    }}
    .premise {{ font-style: italic; color: #4a4338; max-width: 22rem; }}
    .genre {{ font-size: 0.85rem; letter-spacing: 0.12em; text-transform: uppercase;
              color: #7a6f5c; margin-top: 2rem; }}
    .chapter {{ page-break-before: always; }}
    .chapter:first-of-type {{ page-break-before: auto; }}
    .ch-kicker {{
      text-align: center; text-indent: 0 !important; font-size: 0.8rem;
      letter-spacing: 0.18em; text-transform: uppercase; color: #665b4a;
      margin: 3rem 0 0.5rem;
    }}
    .ch-title {{
      text-align: center; font-weight: 600; font-size: 1.35rem;
      margin: 0 0 2.25rem; page-break-after: avoid;
    }}
    p {{ margin: 0; text-indent: 1.5em; text-align: justify; }}
    p.first {{ text-indent: 0; margin-top: 0; }}
    p.scenebreak {{
      text-indent: 0; text-align: center; margin: 1.4em 0;
      letter-spacing: 0.45em; color: #665b4a;
    }}
    @media print {{
      body {{ background: white; margin: 0; max-width: none; padding: 0; }}
    }}
  </style>
</head>
<body>
  <header class="titlepage">
    <h1 class="title">{title}</h1>
    {"".join(meta_bits)}
  </header>
  {"".join(chapters_html)}
</body>
</html>
"""


def _to_docx(project: Project) -> bytes:
    try:
        from docx import Document
        from docx.shared import Pt
    except ImportError as exc:
        raise RuntimeError(
            "python-docx is required for DOCX export. pip install python-docx"
        ) from exc

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    doc.add_heading(project.title or "Untitled", level=0)
    if project.premise:
        p = doc.add_paragraph()
        run = p.add_run(project.premise)
        run.italic = True
    if project.genre:
        doc.add_paragraph(f"Genre: {project.genre}")

    if project.characters:
        doc.add_heading("Dramatis Personae", level=1)
        for c in project.characters:
            line = c.name + (f" — {c.role}" if c.role else "")
            doc.add_paragraph(line, style="List Bullet")

    for i, ch in enumerate(_sorted_chapters(project), start=1):
        doc.add_heading(ch.title or f"Chapter {i}", level=1)
        paras = _paragraphs(ch.content or "")
        if not paras:
            doc.add_paragraph("(empty)")
        else:
            for para in paras:
                doc.add_paragraph(para.replace("\n", " "))

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _story_chapters(project: Project) -> list[list[str]]:
    """Chapter body paragraphs in reading order, per chapter, no titles."""
    out: list[list[str]] = []
    for ch in _sorted_chapters(project):
        out.append(_paragraphs(ch.content or ""))
    return out


def _to_docx_manuscript(project: Project) -> bytes:
    """Manuscript DOCX: story prose only, no titles, no front matter."""
    try:
        from docx import Document
        from docx.shared import Pt
    except ImportError as exc:
        raise RuntimeError(
            "python-docx is required for DOCX export. pip install python-docx"
        ) from exc

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    style.paragraph_format.first_line_indent = Pt(24)
    style.paragraph_format.line_spacing = 2

    from docx.enum.text import WD_ALIGN_PARAGRAPH

    chapters = _story_chapters(project)
    for idx, paras in enumerate(chapters):
        if idx > 0:
            sep = doc.add_paragraph("* * *")
            sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
            sep.paragraph_format.first_line_indent = Pt(0)
        for para in paras:
            doc.add_paragraph(para.replace("\n", " "))

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _manuscript_epub_css() -> str:
    """Minimal stylesheet for the story-only manuscript EPUB."""
    return """
body {
  margin: 0;
  padding: 0;
  font-family: "Palatino Linotype", Palatino, Georgia, "Times New Roman", serif;
  font-size: 1em;
  line-height: 1.5;
  text-align: justify;
  widows: 2;
  orphans: 2;
  hyphens: auto;
  -epub-hyphens: auto;
  -webkit-hyphens: auto;
}
p {
  margin: 0 0 2em 0;
  padding: 0;
  text-indent: 1.5em;
  line-height: 1.5;
}
p.first {
  text-indent: 0;
}
p.chapnum {
  text-indent: 0;
  text-align: center;
  font-size: 1.6em;
  margin: 0 0 2em 0;
  page-break-before: always;
  break-before: page;
  -epub-page-break-before: always;
}
p.chapnum.first {
  page-break-before: auto;
  break-before: auto;
  -epub-page-break-before: auto;
}
""".strip()


def _to_epub_manuscript(project: Project) -> bytes:
    """Manuscript EPUB: story prose only, single xhtml, no cover/TOC/copyright.

    Still a valid EPUB 3.0 with a nav document, so it opens in any reader.
    """
    from datetime import datetime, timezone

    title = project.title or "Untitled"
    lang = (project.language or "en").strip() or "en"
    book_id = f"urn:ghostwriter:{_slug(title)}-manuscript"
    modified = datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ")

    chapters = _story_chapters(project)
    body_parts: list[str] = []
    for idx, paras in enumerate(chapters):
        num = idx + 1
        cls = ' class="chapnum first"' if idx == 0 else ' class="chapnum"'
        body_parts.append(f"<p{cls}>{num}</p>")
        first_p = True
        for p in paras:
            pcls = ' class="first"' if first_p else ""
            body_parts.append(f"<p{pcls}>{_xhtml_escape_text(p)}</p>")
            first_p = False
    body_inner = "\n".join(body_parts)
    if not body_inner:
        body_inner = '<p class="first"><em>(empty)</em></p>'

    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", compression=zipfile.ZIP_DEFLATED) as zf:
        zf.writestr("mimetype", "application/epub+zip", compress_type=zipfile.ZIP_STORED)
        zf.writestr(
            "META-INF/container.xml",
            """<?xml version="1.0" encoding="UTF-8"?>
<container version="1.0" xmlns="urn:oasis:names:tc:opendocument:xmlns:container">
  <rootfiles>
    <rootfile full-path="OEBPS/content.opf" media-type="application/oebps-package+xml"/>
  </rootfiles>
</container>
""",
        )
        zf.writestr("OEBPS/style.css", _manuscript_epub_css() + "\n")

        zf.writestr(
            "OEBPS/manuscript.xhtml",
            _epub_wrap(title, "chapter", body_inner),
        )
        zf.writestr(
            "OEBPS/nav.xhtml",
            f"""<?xml version="1.0" encoding="UTF-8"?>
<!DOCTYPE html>
<html xmlns="http://www.w3.org/1999/xhtml" xmlns:epub="http://www.idpf.org/2007/ops" xml:lang="{lang}" lang="{lang}">
<head>
  <meta charset="utf-8"/>
  <title>Contents</title>
  <link rel="stylesheet" type="text/css" href="style.css"/>
</head>
<body>
  <nav epub:type="toc" id="toc">
    <ol>
      <li><a href="manuscript.xhtml">{_xhtml_escape_text(title)}</a></li>
    </ol>
  </nav>
</body>
</html>
""",
        )
        zf.writestr(
            "OEBPS/content.opf",
            f"""<?xml version="1.0" encoding="UTF-8"?>
<package xmlns="http://www.idpf.org/2007/opf" unique-identifier="BookId" version="3.0">
  <metadata xmlns:dc="http://purl.org/dc/elements/1.1/">
    <dc:title>{_xhtml_escape_text(title)}</dc:title>
    <dc:language>{_xhtml_escape_text(lang)}</dc:language>
    <dc:identifier id="BookId">{_xhtml_escape_text(book_id)}</dc:identifier>
    <meta property="dcterms:modified">{modified}</meta>
  </metadata>
  <manifest>
    <item id="nav" href="nav.xhtml" media-type="application/xhtml+xml" properties="nav"/>
    <item id="css" href="style.css" media-type="text/css"/>
    <item id="manuscript" href="manuscript.xhtml" media-type="application/xhtml+xml"/>
  </manifest>
  <spine page-progression-direction="ltr">
    <itemref idref="manuscript"/>
  </spine>
</package>
""",
        )
        zf.writestr(
            "OEBPS/toc.ncx",
            f"""<?xml version="1.0" encoding="UTF-8"?>
<!DOCTYPE ncx PUBLIC "-//NISO//DTD ncx 2005-1//EN"
  "http://www.daisy.org/z3986/2005/ncx-2005-1.dtd">
<ncx xmlns="http://www.daisy.org/z3986/2005/ncx/" version="2005-1">
  <head>
    <meta name="dtb:uid" content="{_xhtml_escape_text(book_id)}"/>
    <meta name="dtb:depth" content="1"/>
    <meta name="dtb:totalPageCount" content="0"/>
    <meta name="dtb:maxPageNumber" content="0"/>
  </head>
  <docTitle><text>{_xhtml_escape_text(title)}</text></docTitle>
  <navMap>
    <navPoint id="nav1" playOrder="1">
      <navLabel><text>Manuscript</text></navLabel>
      <content src="manuscript.xhtml"/>
    </navPoint>
  </navMap>
</ncx>
""",
        )

    return buf.getvalue()


def _epub_css() -> str:
    """Reflowable novel stylesheet — portrait-friendly, no fixed viewport."""
    return """
/* GhostWriter novel EPUB — reflowable, LTR, reader-controlled type size */
@namespace epub "http://www.idpf.org/2007/ops";

html {
  /* Do not set a fixed width/height/viewport — that forces landscape oddities */
  writing-mode: horizontal-tb;
  direction: ltr;
}

body {
  margin: 0;
  padding: 0;
  font-family: "Palatino Linotype", Palatino, "Book Antiqua", Georgia, "Times New Roman", serif;
  font-size: 1em;              /* respect reader default */
  line-height: 1.5;
  text-align: justify;
  widows: 2;
  orphans: 2;
  hyphens: auto;
  -epub-hyphens: auto;
  -webkit-hyphens: auto;
  adobe-hyphenate: auto;
}

/* Vertical rhythm via em so it scales with reader font size */
p {
  margin: 0;
  padding: 0;
  text-indent: 1.5em;
  line-height: 1.5;
}

p.first {
  text-indent: 0;
}

p.scenebreak {
  text-indent: 0;
  text-align: center;
  margin: 1.25em 0 1.25em 0;
  letter-spacing: 0.55em;
  font-size: 0.95em;
  page-break-inside: avoid;
}

/* —— Title page —— */
body.titlepage {
  text-align: center;
  padding: 30% 1.2em 2em 1.2em;
}

body.titlepage h1 {
  font-size: 1.85em;
  font-weight: normal;
  letter-spacing: 0.06em;
  margin: 0 0 1.75em 0;
  line-height: 1.25;
  text-align: center;
  page-break-after: avoid;
}

body.titlepage .premise {
  font-style: italic;
  font-size: 0.95em;
  text-indent: 0;
  text-align: center;
  margin: 0.75em 1em 0 1em;
  line-height: 1.45;
}

body.titlepage .genre {
  font-style: normal;
  font-size: 0.75em;
  letter-spacing: 0.16em;
  text-transform: uppercase;
  text-indent: 0;
  text-align: center;
  margin-top: 2.5em;
  opacity: 0.75;
}

body.titlepage .byline {
  font-size: 0.9em;
  text-indent: 0;
  text-align: center;
  margin-top: 3em;
  letter-spacing: 0.04em;
}

/* —— Chapters —— */
body.chapter {
  padding: 0 0 1em 0;
}

.ch-open {
  margin: 2.5em 0 2em 0;
  text-align: center;
  page-break-after: avoid;
  -webkit-column-break-after: avoid;
  break-after: avoid;
}

.ch-kicker {
  display: block;
  font-size: 0.8em;
  letter-spacing: 0.22em;
  text-transform: uppercase;
  text-indent: 0;
  text-align: center;
  margin: 0 0 0.6em 0;
  font-weight: normal;
}

.ch-title {
  display: block;
  font-size: 1.35em;
  font-weight: normal;
  letter-spacing: 0.03em;
  text-indent: 0;
  text-align: center;
  margin: 0;
  line-height: 1.3;
}

/* Body copy sits a beat after the opener */
.ch-open + p {
  margin-top: 0.25em;
}

/* TOC page */
body.toc {
  padding: 2em 1.2em;
}
body.toc h1 {
  font-size: 1.2em;
  font-weight: normal;
  letter-spacing: 0.2em;
  text-transform: uppercase;
  text-align: center;
  margin: 0 0 1.75em 0;
}
body.toc ol {
  list-style: none;
  padding: 0;
  margin: 0;
}
body.toc li {
  margin: 0 0 0.85em 0;
  text-align: left;
  line-height: 1.4;
}
body.toc a {
  text-decoration: none;
  color: inherit;
}

/* Cover (SVG wrapper page) */
body.cover {
  margin: 0;
  padding: 0;
  text-align: center;
}
body.cover svg {
  width: 100%;
  height: auto;
  max-height: 100%;
}

/* Copyright page */
body.copy {
  padding: 2em 1.2em;
  text-align: left;
}
body.copy .copy-line {
  text-indent: 0;
  margin: 0 0 1em 0;
  font-size: 0.9em;
  line-height: 1.5;
}
""".strip()


def _epub_wrap(title: str, body_class: str, body_inner: str, css_href: str = "style.css") -> str:
    return f"""<?xml version="1.0" encoding="UTF-8"?>
<!DOCTYPE html PUBLIC "-//W3C//DTD XHTML 1.1//EN"
  "http://www.w3.org/TR/xhtml11/DTD/xhtml11.dtd">
<html xmlns="http://www.w3.org/1999/xhtml" xmlns:epub="http://www.idpf.org/2007/ops" xml:lang="en">
<head>
  <meta http-equiv="Content-Type" content="application/xhtml+xml; charset=utf-8"/>
  <title>{_xhtml_escape_text(title)}</title>
  <link rel="stylesheet" type="text/css" href="{css_href}"/>
</head>
<body class="{body_class}">
{body_inner}
</body>
</html>
"""


def _epub_body_from_blocks(blocks: list[tuple[str, str]]) -> str:
    parts: list[str] = []
    first_p = True
    for kind, text in blocks:
        if kind == "break":
            parts.append('<p class="scenebreak">* * *</p>')
            first_p = True
            continue
        cls = ' class="first"' if first_p else ""
        parts.append(f"<p{cls}>{_xhtml_escape_text(text)}</p>")
        first_p = False
    return "\n".join(parts) if parts else '<p class="first"><em>(empty)</em></p>'


def _epub_cover_svg(title: str, premise: str = "") -> str:
    """Simple portrait cover — 2:3 ratio, dark paper + gold type."""
    t = _xhtml_escape_text(title)
    # Rough wrap for long titles
    words = title.split()
    lines: list[str] = []
    cur = ""
    for w in words:
        trial = f"{cur} {w}".strip()
        if len(trial) > 18 and cur:
            lines.append(cur)
            cur = w
        else:
            cur = trial
    if cur:
        lines.append(cur)
    if not lines:
        lines = [title or "Untitled"]
    lines = lines[:5]
    start_y = 420 - (len(lines) - 1) * 28
    text_nodes = []
    for i, line in enumerate(lines):
        text_nodes.append(
            f'<text x="300" y="{start_y + i * 56}" text-anchor="middle" '
            f'font-family="Georgia, serif" font-size="44" fill="#e8d5a3">{_xhtml_escape_text(line)}</text>'
        )
    sub = ""
    if premise:
        short = premise if len(premise) < 90 else premise[:87] + "…"
        sub = (
            f'<text x="300" y="620" text-anchor="middle" font-family="Georgia, serif" '
            f'font-size="18" font-style="italic" fill="#b8ae9a">{_xhtml_escape_text(short)}</text>'
        )
    return f"""<?xml version="1.0" encoding="UTF-8"?>
<svg xmlns="http://www.w3.org/2000/svg" width="600" height="900" viewBox="0 0 600 900" version="1.1">
  <rect width="600" height="900" fill="#16140f"/>
  <rect x="28" y="28" width="544" height="844" fill="none" stroke="#c4a35a" stroke-width="2" opacity="0.55"/>
  <rect x="40" y="40" width="520" height="820" fill="none" stroke="#c4a35a" stroke-width="0.75" opacity="0.3"/>
  <line x1="120" y1="320" x2="480" y2="320" stroke="#c4a35a" stroke-width="1" opacity="0.4"/>
  {"".join(text_nodes)}
  <line x1="120" y1="560" x2="480" y2="560" stroke="#c4a35a" stroke-width="1" opacity="0.4"/>
  {sub}
  <text x="300" y="820" text-anchor="middle" font-family="Georgia, serif" font-size="14"
        letter-spacing="4" fill="#7f725c">GHOSTWRITER</text>
</svg>
"""


def _cover_font(size: int, italic: bool = False):
    """Load a serif TTF for cover rendering; fall back to default bitmap."""
    from PIL import ImageFont

    candidates = [
        "/usr/share/fonts/truetype/liberation-serif-fonts/"
        + ("LiberationSerif-Italic.ttf" if italic else "LiberationSerif-Regular.ttf"),
        "/usr/share/fonts/truetype/liberation2/LiberationSerif-"
        + ("Italic" if italic else "Regular") + ".ttf",
        "/usr/share/fonts/google-noto-vf/NotoSerif"
        + ("-Italic" if italic else "") + "[wght].ttf",
    ]
    for path in candidates:
        try:
            return ImageFont.truetype(path, size)
        except OSError:
            continue
    return ImageFont.load_default()


def _cover_wrap(title: str, width: int, font) -> list[str]:
    """Greedy word wrap to fit the cover width (rough, letter-based)."""
    words = (title or "Untitled").split()
    lines: list[str] = []
    cur = ""
    for w in words:
        trial = f"{cur} {w}".strip()
        if font.getlength(trial) > width and cur:
            lines.append(cur)
            cur = w
        else:
            cur = trial
    if cur:
        lines.append(cur)
    if not lines:
        lines = ["Untitled"]
    # Shrink wrap so at most 5 lines
    while len(lines) > 5:
        last = lines.pop()
        lines[-1] = f"{lines[-1]} {last}"
    return lines


def _to_cover_image(project: Project, fmt: str) -> bytes:
    """Render a simple print cover (2:3) to JPG or TIFF via Pillow.

    Matches the SVG cover look: dark paper, gold borders, serif title.
    """
    from PIL import Image, ImageDraw

    width, height = 1800, 2700
    bg = (22, 20, 15)          # #16140f dark paper
    gold = (196, 163, 90)      # #c4a35a
    cream = (232, 213, 163)    # #e8d5a3
    muted = (184, 174, 154)    # #b8ae9a
    footer = (127, 114, 92)    # #7f725c

    title = project.title or "Untitled"
    premise = project.premise or project.description or ""
    author = project.author or ""

    img = Image.new("RGB", (width, height), bg)
    draw = ImageDraw.Draw(img)

    # Borders
    draw.rectangle([60, 60, width - 60, height - 60], outline=gold, width=6)
    draw.rectangle([90, 90, width - 90, height - 90], outline=gold, width=2)

    # Title
    title_font = _cover_font(120)
    lines = _cover_wrap(title, int(width * 0.72), title_font)
    line_h = 150
    block_h = line_h * len(lines)
    start_y = (height // 2) - block_h // 2 - 120
    for i, line in enumerate(lines):
        bbox = draw.textbbox((0, 0), line, font=title_font)
        x = (width - (bbox[2] - bbox[0])) / 2 - bbox[0]
        draw.text((x, start_y + i * line_h), line, font=title_font, fill=cream)

    # Rules above/below title
    title_bottom = start_y + block_h
    draw.line([width * 0.2, title_bottom + 90, width * 0.8, title_bottom + 90],
              fill=gold, width=3)

    # Premise (italic subtitle)
    if premise:
        sub = premise if len(premise) < 90 else premise[:87] + "…"
        sub_font = _cover_font(52, italic=True)
        bbox = draw.textbbox((0, 0), sub, font=sub_font)
        x = (width - (bbox[2] - bbox[0])) / 2 - bbox[0]
        draw.text((x, title_bottom + 150), sub, font=sub_font, fill=muted)

    # Author line
    if author:
        auth_font = _cover_font(56)
        bbox = draw.textbbox((0, 0), author, font=auth_font)
        x = (width - (bbox[2] - bbox[0])) / 2 - bbox[0]
        draw.text((x, height * 0.82), author, font=auth_font, fill=cream)

    # Footer
    footer_font = _cover_font(44)
    label = "GHOSTWRITER"
    spaced = " ".join(label)
    bbox = draw.textbbox((0, 0), spaced, font=footer_font)
    x = (width - (bbox[2] - bbox[0])) / 2 - bbox[0]
    draw.text((x, height * 0.93), spaced, font=footer_font, fill=footer)

    buf = io.BytesIO()
    if fmt == "cover-tiff":
        img.save(buf, format="TIFF", dpi=(300, 300))
    else:
        img.save(buf, format="JPEG", quality=92, dpi=(300, 300))
    return buf.getvalue()


def _to_epub(project: Project) -> bytes:
    """Novel-oriented reflowable EPUB 3.0 (with NCX) via stdlib zip.

    Metadata is tuned for Amazon KDP: creator, publisher, rights (copyright),
    ISBN as the identifier when present, series collection, and language.
    """
    from datetime import datetime, timezone

    title = project.title or "Untitled"
    author = project.author or ""
    publisher = project.publisher or "GhostWriter"
    lang = (project.language or "en").strip() or "en"
    if project.isbn:
        book_id = f"urn:isbn:{project.isbn}"
    else:
        book_id = f"urn:ghostwriter:{_slug(title)}"
    chapters = _sorted_chapters(project)
    modified = datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ")
    copyright_year = datetime.now(timezone.utc).year
    first_chap_href = "chap_001.xhtml" if chapters else "title.xhtml"

    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", compression=zipfile.ZIP_DEFLATED) as zf:
        zf.writestr("mimetype", "application/epub+zip", compress_type=zipfile.ZIP_STORED)
        zf.writestr(
            "META-INF/container.xml",
            """<?xml version="1.0" encoding="UTF-8"?>
<container version="1.0" xmlns="urn:oasis:names:tc:opendocument:xmlns:container">
  <rootfiles>
    <rootfile full-path="OEBPS/content.opf" media-type="application/oebps-package+xml"/>
  </rootfiles>
</container>
""",
        )

        zf.writestr("OEBPS/style.css", _epub_css() + "\n")
        cover_svg = _epub_cover_svg(title, project.premise or project.description or "")
        zf.writestr("OEBPS/cover.svg", cover_svg)

        cover_inner = re.sub(r"<\?xml[^?]*\?>\s*", "", cover_svg)
        zf.writestr(
            "OEBPS/cover.xhtml",
            _epub_wrap(f"Cover — {title}", "cover", cover_inner),
        )

        tp_bits = [f"  <h1>{_xhtml_escape_text(title)}</h1>"]
        if author:
            tp_bits.append(
                f'  <p class="byline">{_xhtml_escape_text(author)}</p>'
            )
        if project.premise:
            tp_bits.append(
                f'  <p class="premise">{_xhtml_escape_text(project.premise)}</p>'
            )
        if project.genre:
            tp_bits.append(
                f'  <p class="genre">{_xhtml_escape_text(project.genre)}</p>'
            )
        zf.writestr(
            "OEBPS/title.xhtml",
            _epub_wrap(title, "titlepage", "\n".join(tp_bits)),
        )

        copyright_line = project.copyright or (
            f"Copyright © {copyright_year} {author or title}"
        )
        rights_holder = project.copyright or ""
        copy_bits = [
            f'  <p class="copy-line">{_xhtml_escape_text(copyright_line)}</p>',
            f'  <p class="copy-line">Published by {_xhtml_escape_text(publisher)}</p>',
        ]
        if project.isbn:
            copy_bits.append(
                f'  <p class="copy-line">ISBN: {_xhtml_escape_text(project.isbn)}</p>'
            )
        if project.series:
            pos = f", Book {project.series_position}" if project.series_position else ""
            copy_bits.append(
                f'  <p class="copy-line">{_xhtml_escape_text(project.series)}{_xhtml_escape_text(pos)}</p>'
            )
        copy_bits.append(
            '  <p class="copy-line">All rights reserved. No part of this publication may be '
            "reproduced, distributed, or transmitted in any form or by any means without the "
            "prior written permission of the publisher.</p>"
        )
        zf.writestr(
            "OEBPS/copyright.xhtml",
            _epub_wrap(f"Copyright — {title}", "copy", "\n".join(copy_bits)),
        )

        toc_items = []
        for i, ch in enumerate(chapters, start=1):
            ct = ch.title or f"Chapter {i}"
            toc_items.append(
                f'    <li><a href="chap_{i:03d}.xhtml">{_xhtml_escape_text(ct)}</a></li>'
            )
        zf.writestr(
            "OEBPS/toc.xhtml",
            _epub_wrap(
                "Contents",
                "toc",
                "  <h1>Contents</h1>\n  <ol>\n"
                + ("\n".join(toc_items) if toc_items else "    <li><em>No chapters</em></li>")
                + "\n  </ol>",
            ),
        )

        manifest_items = [
            '<item id="ncx" href="toc.ncx" media-type="application/x-dtbncx+xml"/>',
            '<item id="nav" href="nav.xhtml" media-type="application/xhtml+xml" properties="nav"/>',
            '<item id="css" href="style.css" media-type="text/css"/>',
            '<item id="cover-img" href="cover.svg" media-type="image/svg+xml" properties="cover-image"/>',
            '<item id="cover" href="cover.xhtml" media-type="application/xhtml+xml"/>',
            '<item id="title" href="title.xhtml" media-type="application/xhtml+xml"/>',
            '<item id="copyright" href="copyright.xhtml" media-type="application/xhtml+xml"/>',
            '<item id="tocpage" href="toc.xhtml" media-type="application/xhtml+xml"/>',
        ]
        spine_items = [
            '<itemref idref="cover" linear="no"/>',
            '<itemref idref="title"/>',
            '<itemref idref="copyright"/>',
            '<itemref idref="tocpage"/>',
        ]
        nav_points = [
            """<navPoint id="nav-title" playOrder="1">
  <navLabel><text>Title Page</text></navLabel>
  <content src="title.xhtml"/>
</navPoint>""",
            """<navPoint id="nav-copy" playOrder="2">
  <navLabel><text>Copyright</text></navLabel>
  <content src="copyright.xhtml"/>
</navPoint>""",
            """<navPoint id="nav-toc" playOrder="3">
  <navLabel><text>Contents</text></navLabel>
  <content src="toc.xhtml"/>
</navPoint>""",
        ]
        nav_ol = [
            '      <li><a href="title.xhtml">Title Page</a></li>',
            '      <li><a href="copyright.xhtml">Copyright</a></li>',
            '      <li><a href="toc.xhtml">Contents</a></li>',
        ]
        play = 4

        for i, ch in enumerate(chapters, start=1):
            ct = ch.title or f"Chapter {i}"
            kicker, sub = _chapter_heading_parts(ct, i)
            if sub:
                opener = (
                    f'  <div class="ch-open">\n'
                    f'    <p class="ch-kicker">{_xhtml_escape_text(kicker)}</p>\n'
                    f'    <h1 class="ch-title">{_xhtml_escape_text(sub)}</h1>\n'
                    f"  </div>"
                )
            else:
                opener = (
                    f'  <div class="ch-open">\n'
                    f'    <h1 class="ch-title">{_xhtml_escape_text(kicker)}</h1>\n'
                    f"  </div>"
                )
            body = _epub_body_from_blocks(_novel_blocks(ch.content or ""))
            href = f"chap_{i:03d}.xhtml"
            zf.writestr(
                f"OEBPS/{href}",
                _epub_wrap(ct, "chapter", opener + "\n" + body),
            )
            mid = f"chap{i}"
            manifest_items.append(
                f'<item id="{mid}" href="{href}" media-type="application/xhtml+xml"/>'
            )
            spine_items.append(f'<itemref idref="{mid}"/>')
            nav_points.append(
                f"""<navPoint id="nav{i}" playOrder="{play}">
  <navLabel><text>{_xhtml_escape_text(ct)}</text></navLabel>
  <content src="{href}"/>
</navPoint>"""
            )
            nav_ol.append(
                f'      <li><a href="{href}">{_xhtml_escape_text(ct)}</a></li>'
            )
            play += 1

        zf.writestr(
            "OEBPS/nav.xhtml",
            f"""<?xml version="1.0" encoding="UTF-8"?>
<!DOCTYPE html>
<html xmlns="http://www.w3.org/1999/xhtml" xmlns:epub="http://www.idpf.org/2007/ops" xml:lang="{lang}" lang="{lang}">
<head>
  <meta charset="utf-8"/>
  <title>Navigation</title>
  <link rel="stylesheet" type="text/css" href="style.css"/>
</head>
<body>
  <nav epub:type="toc" id="toc">
    <h1>Contents</h1>
    <ol>
{chr(10).join(nav_ol)}
    </ol>
  </nav>
  <nav epub:type="landmarks" id="landmarks" hidden="hidden">
    <ol>
      <li><a epub:type="cover" href="cover.xhtml">Cover</a></li>
      <li><a epub:type="titlepage" href="title.xhtml">Title Page</a></li>
      <li><a epub:type="toc" href="toc.xhtml">Contents</a></li>
      <li><a epub:type="bodymatter" href="{first_chap_href}">Start</a></li>
    </ol>
  </nav>
</body>
</html>
""",
        )

        desc = project.description or project.premise or ""
        series_meta = ""
        if project.series:
            series_meta = (
                f'    <meta property="belongs-to-collection" id="crt-series">{_xhtml_escape_text(project.series)}</meta>\n'
                '    <meta refines="#crt-series" property="collection-type">series</meta>\n'
                + (
                    f'    <meta refines="#crt-series" property="group-position">{int(project.series_position or 0)}</meta>\n'
                    if project.series_position
                    else ""
                )
            )
        zf.writestr(
            "OEBPS/content.opf",
            f"""<?xml version="1.0" encoding="UTF-8"?>
<package xmlns="http://www.idpf.org/2007/opf" unique-identifier="BookId" version="3.0"
         prefix="rendition: http://www.idpf.org/vocab/rendition/#">
  <metadata xmlns:dc="http://purl.org/dc/elements/1.1/">
    <dc:title>{_xhtml_escape_text(title)}</dc:title>
    {f'<dc:creator id="crt-author">{_xhtml_escape_text(author)}</dc:creator>' if author else ''}
    <dc:language>{_xhtml_escape_text(lang)}</dc:language>
    <dc:identifier id="BookId">{_xhtml_escape_text(book_id)}</dc:identifier>
    <dc:publisher>{_xhtml_escape_text(publisher)}</dc:publisher>
    {f'<dc:rights>{_xhtml_escape_text(copyright_line)}</dc:rights>' if copyright_line else ''}
    <dc:date>{copyright_year}-01-01T00:00:00Z</dc:date>
    {f'<dc:description>{_xhtml_escape_text(desc)}</dc:description>' if desc else ''}
    <meta name="cover" content="cover-img"/>
    <meta property="rendition:layout">reflowable</meta>
    <meta property="rendition:orientation">auto</meta>
    <meta property="rendition:spread">none</meta>
    <meta property="dcterms:modified">{modified}</meta>
{series_meta}  </metadata>
  <manifest>
{chr(10).join("    " + m for m in manifest_items)}
  </manifest>
  <spine toc="ncx" page-progression-direction="ltr">
{chr(10).join("    " + s for s in spine_items)}
  </spine>
  <guide>
    <reference type="cover" title="Cover" href="cover.xhtml"/>
    <reference type="title-page" title="Title Page" href="title.xhtml"/>
    <reference type="copyright-page" title="Copyright" href="copyright.xhtml"/>
    <reference type="toc" title="Contents" href="toc.xhtml"/>
    <reference type="text" title="Start" href="{first_chap_href}"/>
  </guide>
</package>
""",
        )

        zf.writestr(
            "OEBPS/toc.ncx",
            f"""<?xml version="1.0" encoding="UTF-8"?>
<!DOCTYPE ncx PUBLIC "-//NISO//DTD ncx 2005-1//EN"
  "http://www.daisy.org/z3986/2005/ncx-2005-1.dtd">
<ncx xmlns="http://www.daisy.org/z3986/2005/ncx/" version="2005-1">
  <head>
    <meta name="dtb:uid" content="{_xhtml_escape_text(book_id)}"/>
    <meta name="dtb:depth" content="1"/>
    <meta name="dtb:totalPageCount" content="0"/>
    <meta name="dtb:maxPageNumber" content="0"/>
  </head>
  <docTitle><text>{_xhtml_escape_text(title)}</text></docTitle>
  <navMap>
{chr(10).join(nav_points)}
  </navMap>
</ncx>
""",
        )

    return buf.getvalue()


SUPPORTED_FORMATS = (
    "markdown",
    "txt",
    "html",
    "docx",
    "epub",
    "json",
    "manuscript-docx",
    "manuscript-epub",
    "cover-jpg",
    "cover-tiff",
)
